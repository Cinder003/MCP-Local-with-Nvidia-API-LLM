# Enhanced FastMCP Server with comprehensive file creation and system operations
from fastmcp import FastMCP
from fastmcp.exceptions import ToolError
from pathlib import Path
import os
import shutil
import json
import subprocess
import sqlite3
import pandas as pd
import datetime
import psutil
import platform
import glob
import zipfile
import tempfile
import logging
import signal
import sys
from typing import List, Dict, Any, Optional, Union
import openpyxl
from docx import Document
import csv
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Load configuration
def load_config():
    """Load configuration from JSON file with environment variable fallbacks"""
    config_path = 'server_config.json'
    if os.path.exists(config_path):
        with open(config_path, 'r') as f:
            config = json.load(f)
    else:
        # Default configuration if file not found
        config = {
            "server": {
                "name": "Enhanced Comprehensive System Orchestration Server",
                "system_wide_access": True,
                "max_file_size": 104857600,
                "allowed_extensions": [".txt", ".csv", ".xlsx", ".docx", ".json"],
                "restricted_paths": ["C:\\Windows\\System32", "/bin", "/sbin"]
            },
            "logging": {"level": "INFO"},
            "shell": {"timeout": 30},
            "file_operations": {"default_encoding": "utf-8", "max_lines_preview": 1000}
        }

    # Override with environment variables if present
    config["server"]["system_wide_access"] = os.getenv("SYSTEM_WIDE_ACCESS", str(config["server"]["system_wide_access"])).lower() == "true"
    config["server"]["max_file_size"] = int(os.getenv("MAX_FILE_SIZE", config["server"]["max_file_size"]))
    config["logging"]["level"] = os.getenv("LOG_LEVEL", config["logging"]["level"])
    config["shell"]["timeout"] = int(os.getenv("SHELL_TIMEOUT", config["shell"]["timeout"]))

    return config

# Load configuration
CONFIG = load_config()

# Configure logging
logging.basicConfig(level=getattr(logging, CONFIG["logging"]["level"]))
logger = logging.getLogger(__name__)

# Create the FastMCP server instance
mcp = FastMCP(
    name=CONFIG["server"]["name"],
    instructions=f"""
This server provides comprehensive system orchestration capabilities including:

- Universal file creation (TXT, CSV, XLSX, DOCX, JSON, etc.)
- File and directory management (create, copy, move, delete, search, read, edit)
- System operations (shell commands, process management, app launching)
- Data processing (format conversion, extraction, database operations)
- Archive operations (zip creation/extraction with proper folder structure preservation)

Configuration loaded from: server_config.json and environment variables
System-wide access: {CONFIG["server"]["system_wide_access"]}
Max file size: {CONFIG["server"]["max_file_size"]} bytes
"""
)

# Configuration class for easy access
class ServerConfig:
    SYSTEM_WIDE_ACCESS = CONFIG["server"]["system_wide_access"]
    MAX_FILE_SIZE = CONFIG["server"]["max_file_size"]
    ALLOWED_EXTENSIONS = CONFIG["server"]["allowed_extensions"]
    RESTRICTED_PATHS = CONFIG["server"]["restricted_paths"]
    DEFAULT_ENCODING = CONFIG["file_operations"]["default_encoding"]
    MAX_LINES_PREVIEW = CONFIG["file_operations"]["max_lines_preview"]
    SHELL_TIMEOUT = CONFIG["shell"]["timeout"]

config = ServerConfig()

def is_path_safe(path: str) -> bool:
    """Enhanced path validation for system-wide access"""
    if not config.SYSTEM_WIDE_ACCESS:
        return True

    try:
        abs_path = os.path.abspath(path)
        for restricted in config.RESTRICTED_PATHS:
            if abs_path.startswith(restricted):
                return False
        return True
    except Exception:
        return False

# [Rest of the server tools remain the same - they will use config.* instead of hardcoded values]
# ... [Include all the existing tools with config references]

def format_file_size(size_bytes: int) -> str:
    """Format file size in human-readable format"""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024.0:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024.0
    return f"{size_bytes:.1f} TB"

# ENHANCED FILE CREATION TOOLS
@mcp.tool
def create_file(
    path: str,
    content: str = "",
    file_type: str = "auto",
    working_directory: str = None,
    encoding: str = None
) -> str:
    """Create any type of file with appropriate content structure."""
    encoding = encoding or config.DEFAULT_ENCODING

    try:
        if working_directory:
            full_path = os.path.join(working_directory, path) if not os.path.isabs(path) else path
        else:
            full_path = path

        if not is_path_safe(full_path):
            raise ToolError("Access denied: Path is in a restricted system directory")

        # Create parent directories if needed
        parent_dir = os.path.dirname(full_path)
        if parent_dir:
            os.makedirs(parent_dir, exist_ok=True)

        # Detect file type from extension if auto
        ext = os.path.splitext(full_path)[1].lower()
        if file_type == "auto":
            file_type = ext.lstrip('.')

        # Check file size limit
        if len(content.encode(encoding)) > config.MAX_FILE_SIZE:
            raise ToolError(f"Content exceeds maximum file size of {format_file_size(config.MAX_FILE_SIZE)}")

        # Create file based on type
        if file_type in ['txt', 'md', 'log', 'py', 'js', 'html', 'css', 'json', 'xml', 'yaml', 'sql', 'sh', 'bat']:
            # Text-based files
            with open(full_path, 'w', encoding=encoding) as f:
                f.write(content)
        elif file_type == 'csv':
            # CSV file creation
            with open(full_path, 'w', newline='', encoding=encoding) as f:
                if content:
                    f.write(content)
                else:
                    # Create sample CSV structure
                    writer = csv.writer(f)
                    writer.writerow(['Column1', 'Column2', 'Column3'])
                    writer.writerow(['Value1', 'Value2', 'Value3'])
        elif file_type in ['xlsx', 'xls']:
            # Excel file creation
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Sheet1"

            if content:
                # Try to parse content as CSV-like data
                lines = content.strip().split('\n')
                for row_idx, line in enumerate(lines, 1):
                    values = line.split(',')
                    for col_idx, value in enumerate(values, 1):
                        ws.cell(row=row_idx, column=col_idx, value=value.strip())
            else:
                # Create sample Excel structure
                ws['A1'] = 'Header1'
                ws['B1'] = 'Header2' 
                ws['C1'] = 'Header3'
                ws['A2'] = 'Data1'
                ws['B2'] = 'Data2'
                ws['C2'] = 'Data3'

            wb.save(full_path)
        elif file_type in ['docx', 'doc']:
            # Word document creation
            doc = Document()
            if content:
                doc.add_paragraph(content)
            else:
                doc.add_heading('Document Title', 0)
                doc.add_paragraph('This is a sample paragraph.')
            doc.save(full_path)
        else:
            # Default to text file
            with open(full_path, 'w', encoding=encoding) as f:
                f.write(content)

        file_info = os.stat(full_path)
        abs_path = os.path.abspath(full_path)

        return f"""File created successfully!
Path: {abs_path}
Type: {file_type.upper()}
Size: {format_file_size(file_info.st_size)}
Content length: {len(content)} characters"""

    except Exception as e:
        raise ToolError(f"Failed to create file: {str(e)}")

# [Continue with all other tools using config references...]
# For brevity, showing the pattern - all tools would be updated similarly

if __name__ == "__main__":
    # Create workspace directory if it doesn't exist
    Path("workspace").mkdir(exist_ok=True)

    print("Starting ENHANCED Comprehensive FastMCP Server...")
    print("✓ Configuration loaded from server_config.json and environment variables")
    print(f"✓ System-wide access: {config.SYSTEM_WIDE_ACCESS}")
    print(f"✓ Max file size: {format_file_size(config.MAX_FILE_SIZE)}")
    print(f"✓ Logging level: {CONFIG['logging']['level']}")

    try:
        # Register signal handlers
        signal.signal(signal.SIGINT, lambda s, f: sys.exit(0))
        signal.signal(signal.SIGTERM, lambda s, f: sys.exit(0))

        # Run FastMCP server with proper shutdown handling
        mcp.run(transport="stdio")
    except KeyboardInterrupt:
        pass  # Clean exit - no print as stdout is closed
    except Exception as e:
        logger.error(f"Server error: {e}")
